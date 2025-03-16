import json
import os
import fitz  # PyMuPDF para PDFs
from docx import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LCDocument
from gerenciador_perguntas import buscar_resposta_sugerida, salvar_pergunta_nao_respondida, resposta_valida

# Diretórios principais
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOADS_DIR = os.path.join(BASE_DIR, "uploads")
TXT_DIR = os.path.join(BASE_DIR, "txt_files")
LOG_DIR = os.path.join(BASE_DIR, "logs")

# Criar diretórios se não existirem
os.makedirs(TXT_DIR, exist_ok=True)
os.makedirs(LOG_DIR, exist_ok=True)

# 📂 Processamento de documentos para geração de embeddings
def extrair_texto(caminho_arquivo):
    """Extrai texto de arquivos PDF, DOCX e TXT."""
    _, ext = os.path.splitext(caminho_arquivo)
    texto = ""

    if ext.lower() == ".pdf":
        with fitz.open(caminho_arquivo) as doc:
            texto = "\n".join([page.get_text("text") for page in doc])

    elif ext.lower() == ".docx":
        doc = Document(caminho_arquivo)
        texto = "\n".join([p.text for p in doc.paragraphs])

    elif ext.lower() == ".txt":
        with open(caminho_arquivo, "r", encoding="utf-8") as f:
            texto = f.read()

    return texto

def processar_arquivos():
    """Lê arquivos da pasta 'uploads', processa e salva como TXT na pasta 'txt_files'."""
    arquivos = [f for f in os.listdir(UPLOADS_DIR) if f.endswith(('.pdf', '.docx', '.txt'))]

    if not arquivos:
        print("⚠️ Nenhum arquivo encontrado para processar.")
        return

    print("📂 Processando arquivos...")

    for arquivo in arquivos:
        caminho_origem = os.path.join(UPLOADS_DIR, arquivo)
        texto_extraido = extrair_texto(caminho_origem)

        if texto_extraido:
            caminho_txt = os.path.join(TXT_DIR, f"{os.path.splitext(arquivo)[0]}.txt")
            with open(caminho_txt, "w", encoding="utf-8") as f:
                f.write(texto_extraido)
            print(f"✅ Arquivo processado e salvo: {caminho_txt}")

# 🔄 Criando a IA e o retriever com FAISS
def processar_documentos():
    """Lê arquivos TXT e converte em objetos de documentos para geração de embeddings."""
    documentos = []
    arquivos = [f for f in os.listdir(TXT_DIR) if f.endswith(".txt")]

    for arquivo in arquivos:
        caminho = os.path.join(TXT_DIR, arquivo)
        with open(caminho, "r", encoding="utf-8") as f:
            documentos.append(LCDocument(page_content=f.read(), metadata={"source": arquivo}))

    return documentos

# 🚀 Iniciar processamento de arquivos
processar_arquivos()

# Criando embeddings e estrutura de busca
documentos = processar_documentos()
if not documentos:
    raise ValueError("❌ Nenhum documento processado. Verifique os arquivos na pasta 'uploads'.")

splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
docs_splitted = splitter.split_documents(documentos)

print("🔄 Criando embeddings para os documentos...")
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
vectorstore = FAISS.from_documents(docs_splitted, embeddings)

# Criando a IA e a estrutura de recuperação de contexto
chain = ConversationalRetrievalChain.from_llm(
    llm=ChatOpenAI(model="gpt-3.5-turbo", temperature=0, api_key=os.getenv("OPENAI_API_KEY")),
    retriever=vectorstore.as_retriever(search_kwargs={"k": 5}),
    return_source_documents=True
)

# 🔄 Loop principal do chatbot
historico = []
while True:
    pergunta = input("\nFaça uma pergunta (ou 'sair' para finalizar'): ")
    if pergunta.lower() == 'sair':
        print("👋 Até logo!")
        break

    # 🔹 Primeiro, verifica se há uma resposta salva manualmente
    resposta_sugerida = buscar_resposta_sugerida(pergunta)
    if resposta_sugerida:
        print("\n🤖 Resposta (Sugestão Manual):", resposta_sugerida)
        continue

    # 🔹 Se não encontrou, busca a resposta gerada pela IA
    resposta = chain.invoke({"question": pergunta, "chat_history": historico})
    resposta_texto = resposta["answer"]

    # 🔹 Se a resposta não for válida, salva no JSON de perguntas não respondidas
    if not resposta_valida(resposta_texto):
        print("\n🤖", resposta_texto)  # Exibe a resposta genérica
        salvar_pergunta_nao_respondida(pergunta, resposta_texto)
    else:
        print("\n🤖 Resposta:", resposta_texto)

        # 🔹 Exibir fontes utilizadas, se houver
        fontes = set(doc.metadata['source'] for doc in resposta['source_documents'])
        if fontes:
            print("\n📌 Fontes utilizadas:")
            for fonte in fontes:
                print(f"- {fonte}")

    historico.append((pergunta, resposta_texto))  # Adiciona ao histórico